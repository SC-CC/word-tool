# --trusted-host pypi.org --trusted-host files.pythonhosted.org
import os
import io
import re
import requests
import docx
from docx import Document
from docx.oxml.ns import qn
import datetime
from datetime import datetime as dt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt, RGBColor
from tkinter import *
from tkinter import ttk
from threading import Thread as process
import ssl, urllib
import urllib.request
import requests
from bs4 import BeautifulSoup
import selenium
from selenium import webdriver
from selenium.webdriver import ActionChains

from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By

from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import Select
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.chrome.options import Options
from tkinter import messagebox

from time import sleep

boan_count = 1
List_Boan_selecting = []
List_Boan_selected = []
mail_count = 1
LIST_KRCERT = []
KRCERT_selected = []


def Combo_Update():
    def write_for_mail_insert(*args):
        global boan_count
        global List_Boan_selecting
        global List_Boan_selected
        if boan_count == 1:
            TEXT_for_mail.insert(END, "1. 보안기사\n1-{0}. {1}".format(boan_count, CB_for_boan.get()))
        else:
            TEXT_for_mail.insert(END, "\n1-{0}. {1}".format(boan_count, CB_for_boan.get()))
        for list in List_Boan_selecting:
            if CB_for_boan.get() == list[0]:
                List_Boan_selected.insert(100, [list[0], list[1]])
        boan_count = boan_count + 1

    CB_for_boan = ttk.Combobox(root, width=60, height=20, takefocus=NO, values=LIST_Boan)
    CB_for_boan.bind("<<ComboboxSelected>>", write_for_mail_insert)
    CB_for_boan.grid(column=1, row=0, sticky=W + N + E + S, pady=20)


def boannews_crawl():
    global List_Boan_selected
    global List_Boan_selecting
    List_Boan_selected.clear()
    List_Boan_selecting = []

    global boan_count
    boan_count = 1
    TEXT_for_mail.delete(1.0, END)
    TEXT_for_mail.insert(END, " 일일 외부 동향 {}".format(
        dt.strftime(dt.now(), "%Y%m%d")))
    global LIST_Boan
    url = "https://www.boannews.com/media/list.asp?mkind=1"
    soup = BeautifulSoup(urllib.request.urlopen(url).read(), features="html.parser")
    mainList = soup.findAll("div", {"class": "news_main_title"})  # 메인뉴스

    for name in mainList:
        abc = name.findAll("a")
        for a in abc:
            main_href = "https://www.boannews.com{}".format(a.attrs['href'])
            main_text = "{} - 보안뉴스".format(a.string.strip())
            print("Main:\n", main_text, "\n", main_href, "\n")
            LIST_Boan.insert(100, main_text)
            List_Boan_selecting.insert(0, [main_text, main_href])

    # 메인 뉴스가 아닌 목록 추가하기 위함
    nameList1 = soup.findAll("a", {"class": "news_content"})  # 전체
    nameList2 = soup.findAll("span", {"class": "news_txt"})  # 뉴스 제목
    nameList3 = soup.findAll("span", {"class": "news_writer"})  # 날짜
    a = 0
    for i in nameList3:
        print("ㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡ")
        print(nameList2[a].get_text().replace("\n", ""))  # 제목
        sub_day = nameList3[a].get_text()
        sub_day = sub_day[sub_day.find("|") + 2:-5]
        print(sub_day)
        sub_day = sub_day.replace("년 ", ".").replace("월 ", ".").replace("일 ", "")
        sub_day3 = dt.strptime(sub_day, '%Y.%m.%d')
        days_diff = (dt.now() - sub_day3).days
        print(days_diff)  # 기자 | 날짜
        if days_diff < 2:
            print("https://www.boannews.com" + nameList1[a].attrs['href'] + "\n")  # URL
            sub_text = (nameList2[a].get_text().replace("\n", "")).strip() + " - 보안뉴스"
            sub_href = "https://www.boannews.com" + nameList1[a].attrs['href']
            LIST_Boan.insert(100, sub_text)
            List_Boan_selecting.insert(0, [sub_text, sub_href])
        a = a + 1
    Combo_Update()


def EXPLOITDB_Write_text():
    No_count = 50828  # ExploitDB 확인을 시작하는 게시물 번호
    count = 1
    TEXT_for_mail.insert(END, "\n4. 공개된 신규 취약점 (ExploitDB)")
    for i in range(1, 1000):
        try:
            url = "https://www.exploit-db.com/exploits/{0}".format(No_count)
            soup = BeautifulSoup(urllib.request.urlopen(url).read(), features="html.parser")
            code = soup.findAll("h1", {"class": "card-title text-secondary text-center"})
            if code == []:
                print("+1")

                No_count = No_count + 1
                continue
            TITLE = code[0].text.strip()  # 제목
            code = soup.findAll("h6", {"class": "stats-title"})
            DATE = code[5].text.strip()  # 날짜
            TYPE = code[3].text.strip()  # 타입
            days_diff = (dt.now() - dt.strptime(DATE, '%Y-%m-%d')).days
            if days_diff < 8:
                if TYPE == "webapps" or TYPE == "remote" or TYPE == "local":
                    NO = code[0].text.strip()  # Number
                    URL = "https://www.exploit-db.com/exploits/{}".format(code[0].text.strip())  # URL
                    CVE = code[1].text.strip()  # CVE
                    if CVE != "N/A":
                        CVE = "CVE-" + CVE
                    else:
                        CVE = "-"

                    print("번호 : " + NO)
                    print("제목 : " + TITLE)
                    print("타입 : " + TYPE)
                    print("CVE ID : " + CVE)
                    print("date : " + DATE)
                    print("URL : " + URL)
                    print("ㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡ")
                    TEXT_for_mail.insert(END,
                                         "\n4-{0}. {1} - ExploitDB\n- CVE ID : {2}\n- 확인내용 : \n".format(count, TITLE,
                                                                                                        CVE))
                    count = count + 1

                    document.add_paragraph("{} - ExploitDB".format(TITLE), style=boan_style)
                    table = document.add_table(rows=1, cols=1)
                    table.style = document.styles['Table Grid']
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].paragraphs[0].add_run("CVE-ID").bold = True
                    hdr_cells[0].paragraphs[0].add_run("\n" + CVE + "\n\n")
                    hdr_cells[0].paragraphs[0].add_run("Description").bold = True
                    hdr_cells[0].paragraphs[0].add_run("\n" + TITLE + "\n\n")
                    hdr_cells[0].paragraphs[0].add_run("Date Entry Created").bold = True
                    hdr_cells[0].paragraphs[0].add_run("\n" + DATE + "\n\n")
                    hdr_cells[0].paragraphs[0].add_run("POC").bold = True
                    hdr_cells[0].paragraphs[0].add_run("\n" + URL)
                    document.add_paragraph("", style=boan_style)
            No_count = No_count + 1
        except:
            print(No_count - 1)
            print("이게 마지막입니다.")
            break
    if count == 1:
        TEXT_for_mail.insert(END, "\n- 없음\n")
        document.add_paragraph("없음\n", style=boan_style)
        document.add_page_break()


def KRCERT_Write_text():
    global mail_count
    url = "https://www.krcert.or.kr/data/secNoticeList.do"
    soup = BeautifulSoup(urllib.request.urlopen(url).read(), features="html.parser")
    mainList2 = soup.findAll("table", {"class": "basicList default"})

    url = "https://www.krcert.or.kr/data/secInfoList.do"
    soup = BeautifulSoup(urllib.request.urlopen(url).read(), features="html.parser")
    mainList3 = soup.findAll("table", {"class": "basicList default"})

    # 보안공지 부분
    mail_count = 1
    TEXT_for_mail.insert(END, "\n\n2. 보안권고문 및 동향\n")
    i = 2
    for name in mainList2:
        while i < 30:
            for list1 in name.findAll('td', {"class": "colTit"}):
                for list2 in list1.findAll('a'):
                    up_time = name.findAll('td', {"class": "gray"})[i].string
                    up_time2 = dt.strptime(up_time, '%Y.%m.%d')
                    days_diff = (dt.now() - up_time2).days
                    if days_diff < 2:
                        subject_KR1 = "{} - KRCERT(보안공지)".format(list2.string)
                        TEXT_for_mail.insert(END, "2-{0}. {1}\n".format(mail_count, subject_KR1))
                        mail_count = mail_count + 1
                        KRCERT_URL = "https://www.krcert.or.kr{}".format(list2.attrs['href'])
                        soup = BeautifulSoup(urllib.request.urlopen(KRCERT_URL).read(), features="html.parser")
                        test = soup.findAll('td', {"class": "cont"})
                        gogotest = "{}".format(test[0].get_text())
                        gogotest = gogotest.split("\n")
                        document.add_paragraph(subject_KR1, style=boan_style)
                        table = document.add_table(rows=1, cols=1)
                        document.add_paragraph("\n")
                        table.style = document.styles['Table Grid']
                        hdr_cells = table.rows[0].cells
                        for each_line in gogotest:
                            if each_line.find("<td ") != -1 or each_line == "페이스북" or each_line == "트위터" != -1:
                                None
                            elif each_line.find("□ 개 요") != -1 or each_line.find("□ 개요") != -1:
                                hdr_cells[0].paragraphs[0].add_run(each_line).bold = True
                            elif each_line.find("□") != -1:
                                hdr_cells[0].paragraphs[0].add_run("\n" + each_line).bold = True
                            elif each_line.strip() == "":
                                None
                            else:
                                hdr_cells[0].paragraphs[0].add_run(each_line.strip() + "\n")
                        hdr_cells[0].paragraphs[0].add_run("\n\n원문 : {0}".format(KRCERT_URL)).bold = False
                        print(list2.string)
                        print("https://www.krcert.or.kr{}".format(list2.attrs['href']))
                        print('------------------')
                        i = i + 3
                    else:
                        i = i + 3
                        # 취약점 정보 부분
    i = 2
    for name in mainList3:
        while i < 30:
            for list1 in name.findAll('td', {"class": "colTit"}):
                for list2 in list1.findAll('a'):
                    up_time = name.findAll('td', {"class": "gray"})[i].string
                    up_time2 = dt.strptime(up_time, '%Y.%m.%d')
                    days_diff = (dt.now() - up_time2).days
                    if days_diff < 2:
                        # 시간차 1일 이하 만 보고서 작성하기 위함
                        subject_KR2 = "{} - KRCERT(취약점 정보)".format(list2.string)

                        TEXT_for_mail.insert(END, "2-{0}. {1}\n".format(mail_count, subject_KR2))
                        mail_count = mail_count + 1

                        KRCERT_URL = "https://www.krcert.or.kr{}".format(list2.attrs['href'])
                        soup = BeautifulSoup(urllib.request.urlopen(KRCERT_URL).read(), features="html.parser")
                        test = soup.findAll('td', {"class": "cont"})
                        gogotest = "{}".format(test[0].get_text())
                        gogotest = gogotest.split("\n")
                        document.add_paragraph(subject_KR2, style=boan_style)
                        table = document.add_table(rows=1, cols=1)
                        document.add_paragraph("\n")
                        table.style = document.styles['Table Grid']
                        hdr_cells = table.rows[0].cells
                        for each_line in gogotest:
                            if each_line.find("<td ") != -1 or each_line == "페이스북" or each_line == "트위터" != -1:
                                None
                            elif each_line.find("□ 개 요") != -1 or each_line.find("□ 개요") != -1:
                                hdr_cells[0].paragraphs[0].add_run(each_line).bold = True
                            elif each_line.find("□") != -1:
                                hdr_cells[0].paragraphs[0].add_run("\n" + each_line).bold = True
                            elif each_line.strip() == "":
                                None
                            else:
                                hdr_cells[0].paragraphs[0].add_run(each_line.strip() + "\n")
                        hdr_cells[0].paragraphs[0].add_run("\n\n원문 : {0}".format(KRCERT_URL)).bold = False
                        print(list2.string)
                        print("https://www.krcert.or.kr{}".format(list2.attrs['href']))
                        print('------------------')
                        i = i + 3
                    else:
                        i = i + 3


def USCERT_Write_text():
    global mail_count
    url = "https://www.kb.cert.org/vuls/bypublished/desc/"
    soup = BeautifulSoup(urllib.request.urlopen(url).read(), features="html.parser")
    # print(soup)

    soup = soup.findAll("table", {"class": "searchby unstriped scroll"})
    i = 1
    for list in soup:
        for list2 in list.findAll("tbody"):
            for list3 in list.findAll("td"):
                if i % 6 == 1:  # 날짜 추출 용
                    DATE = list3.text.strip()
                    diff = (dt.now() - dt.strptime(DATE, '%Y-%m-%d')).days
                    if diff > 2:
                        break
                    print("ㅡㅡㅡㅡㅡㅡ")
                    print("날짜 : " + DATE)
                elif i % 6 == 0:  # 제목과 제목의 a 태그 URL 추출용
                    TITLE = list3.text.strip() + " - USCERT"
                    URL = "https://www.kb.cert.org{0}".format(list3.findAll("a")[0].attrs["href"])
                    print("제목 : " + TITLE)
                    print("URL : " + URL)
                    print("ㅡㅡㅡㅡㅡㅡ")

                    TEXT_for_mail.insert(END, "2-{0}. {1}\n".format(mail_count, TITLE))
                    document.add_paragraph(TITLE, style=boan_style)
                    table = document.add_table(rows=1, cols=1)
                    table.style = document.styles['Table Grid']
                    hdr_cells = table.rows[0].cells

                    html = requests.get(URL, verify=False)
                    soup = BeautifulSoup(html.text, "html.parser")

                    find_day = soup.findAll("div", {"class": "vulcontent"})
                    pcre = re.compile("20\S*\s\S*\sUTC")
                    test = pcre.findall(find_day[0].text.strip())[0].replace(" UTC", "")
                    today = dt.now()
                    min_time = dt.strptime("{0}-{1}-{2} 15:00".format(today.year, today.month, today.day),
                                           '%Y-%m-%d %H:%M') - datetime.timedelta(days=2)
                    max_time = dt.strptime("{0}-{1}-{2} 14:59".format(today.year, today.month, today.day),
                                           '%Y-%m-%d %H:%M') - datetime.timedelta(days=1)
                    upload_time = dt.strptime(test, '%Y-%m-%d %H:%M')
                    print("USCERT Last_modifed(UTC기준): {}".format(upload_time))
                    print("USCERT Last_modifed(한국기준): {}".format(upload_time + datetime.timedelta(hours=9)))

                    if min_time <= upload_time and max_time >= upload_time:  # UTC기준으로 전전일 15시 ~ 전일 15시 이전 까지가 한국 전일임
                        soup2 = soup.findAll("div", {"class": "large-12 columns"})
                        text = soup2[1].text.strip()
                        for each_line in text.split('\n'):
                            if each_line == "Overview":
                                hdr_cells[0].paragraphs[0].add_run(each_line + "\n").bold = True
                            elif each_line == "Description" or each_line == "Impact" or each_line == "Solution":
                                hdr_cells[0].paragraphs[0].add_run("\n" + each_line + "\n").bold = True
                            elif each_line == "Acknowledgements":
                                break
                            else:
                                hdr_cells[0].paragraphs[0].add_run(each_line + "\n")
                        hdr_cells[0].paragraphs[0].add_run("\n원문 : {0}".format(URL))
                        mail_count = mail_count + 1
                i = i + 1
    if mail_count == 1:
        TEXT_for_mail.insert(END, "- 없음\n")
        document.add_paragraph("없음\n", style=boan_style)


def boannews_Write_text(main_text, main_href):
    soup2 = BeautifulSoup(urllib.request.urlopen(main_href).read(), features="html.parser")
    image_caption = soup2.findAll("p", {"class": "txt"})

    test = soup2.findAll("div", {"id": "news_content"})
    test = "{}".format(test[0]).replace("<br/>", "\n")

    test = test.split("\n")
    document.add_paragraph(main_text, style=boan_style)
    table = document.add_table(rows=1, cols=1)
    table.style = document.styles['Table Grid']
    hdr_cells = table.rows[0].cells
    for each_line in test:
        c = 0
        each_line = "{}\n".format(each_line)
        if each_line[0:3] == "<b>" or each_line[0:3] == "<b " or each_line[-10:].find("</b>") != -1:
            hdr_cells[0].paragraphs[0].add_run(
                each_line.replace("<b>", "").replace("</b>", "").replace("<mark>", "").replace("</mark>", "").replace(
                    '<b style="font-weight:400;color:#000000;">', "").strip() + "\n").bold = True
        elif each_line[:].find("<img") != -1:
            print(each_line[:])
            pcre = re.compile('h\S*.JPG|h\S*.jpg|h\S*.png|h\S*.PNG')
            urls = pcre.findall(each_line)
            response = requests.get(urls[0], stream=True)
            image = io.BytesIO(response.content)
            tables = hdr_cells[0].paragraphs[0]
            tabless = tables.add_run()
            tabless.add_picture(image, width=Cm(15.27))
            try:
                tables.add_run(image_caption[c])
            except:
                None
            tables.add_run("\n\n")
            c = c + 1
        elif each_line[:].find("div>") != -1 or each_line[:].find("<a href") != -1 or each_line[:].find("<div") != -1:
            continue
        elif each_line[:].find("<u>") != -1:
            hdr_cells[0].paragraphs[0].add_run(
                each_line.replace("<u>", "").replace("</u>", "").strip() + "\n").underline = True
        else:
            hdr_cells[0].paragraphs[0].add_run(
                each_line.replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">").strip() + "\n").bold = False

    hdr_cells[0].paragraphs[0].add_run("\n원문 : {0}".format(main_href)).bold = False

    document.add_page_break()


def default_write():
    global document
    dayday = "{}".format(dt.strftime(dt.now(), "%Y.%m.%d"))
    document = docx.Document()
    main = document.styles['Normal']
    main.font.name = '맑은 고딕'
    main.font.size = Pt(12)
    rFonts = main.element.rPr.rFonts
    rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    Head = document.styles.add_style('Head', WD_STYLE_TYPE.PARAGRAPH)
    Head.font.name = '맑은 고딕'
    Head.font.size = Pt(42)
    Head.font.color.rgb = RGBColor(31, 71, 137)
    rFonts2 = Head.element.rPr.rFonts
    rFonts2.set(qn("w:eastAsia"), "맑은 고딕")

    sub_Head = document.styles.add_style('sub_Head', WD_STYLE_TYPE.PARAGRAPH)
    sub_Head.font.name = '맑은 고딕'
    sub_Head.font.size = Pt(12)
    sub_Head.font.color.rgb = RGBColor(31, 71, 137)
    rFonts3 = sub_Head.element.rPr.rFonts
    rFonts3.set(qn("w:eastAsia"), "맑은 고딕")

    sub_Head2 = document.styles.add_style('sub_Head2', WD_STYLE_TYPE.PARAGRAPH)
    sub_Head2.font.name = '맑은 고딕'
    sub_Head2.font.size = Pt(16)
    sub_Head2.font.bold = True
    sub_Head2.font.color.rgb = RGBColor(31, 71, 137)
    rFonts4 = sub_Head2.element.rPr.rFonts
    rFonts4.set(qn("w:eastAsia"), "맑은 고딕")

    boan_style = document.styles.add_style('boan_style', WD_STYLE_TYPE.PARAGRAPH)
    boan_style.font.name = '맑은 고딕'
    boan_style.font.size = Pt(14)
    boan_style.font.bold = True
    rFonts4 = boan_style.element.rPr.rFonts
    rFonts4.set(qn("w:eastAsia"), "맑은 고딕")

    source_url = document.styles.add_style('source_url', WD_STYLE_TYPE.PARAGRAPH)
    source_url.font.name = '맑은 고딕'
    source_url.font.size = Pt(11)
    source_url.font.bold = False
    rFonts4 = boan_style.element.rPr.rFonts
    rFonts4.set(qn("w:eastAsia"), "맑은 고딕")

    document.add_paragraph('\n일일 외부 동향\n', style=Head)

    document.add_paragraph(
        '1. 보안 기사\n\n2. 보안 권고문 및 동향\n\n3. 보안 취약점\n\n4. 공개된 신규 취약점 (ExploitDB)\n\n5. 공개된 신규 취약점 (GitHub)\n\n\n',
        style=sub_Head)
    document.add_paragraph(dayday, style=sub_Head)
    document.add_page_break()

    document.add_paragraph('1. 보안 기사', style=sub_Head2)

    global List_Boan_selected
    print(List_Boan_selected)
    print("확인해봐라")
    for list in List_Boan_selected:
        boannews_Write_text(list[0], list[1])

    document.add_paragraph('2. 보안 권고문 및 동향', style=sub_Head2)

    KRCERT_Write_text()
    USCERT_Write_text()

    document.add_paragraph('3. 보안 취약점', style=sub_Head2)
    document.add_paragraph('금융ISAC 크롤링 불가, 추후 API 배포 시 구현 가능\n', style=boan_style)
    TEXT_for_mail.insert(END, "\n3. 보안 취약점\n3-1. - 금융ISAC\n- CVE ID : \n- 확인내용 :\n")

    document.add_paragraph('4. 공개된 신규 취약점 (ExploitDB)', style=sub_Head2)
    EXPLOITDB_Write_text()

    document.add_page_break()
    document.add_paragraph('5. 공개된 신규 취약점 (GitHub)', style=sub_Head2)
    TEXT_for_mail.insert(END, "\n\n5. 공개된 신규 취약점 (GitHub)\n\n\n감사합니다.")
    document.add_paragraph("있을 경우 - Github", style=boan_style)
    table = document.add_table(rows=1, cols=1)
    table.style = document.styles['Table Grid']
    hdr_cells = table.rows[0].cells
    hdr_cells[0].paragraphs[0].add_run("CVE-ID").bold = True
    hdr_cells[0].paragraphs[0].add_run("\n\n\n")
    hdr_cells[0].paragraphs[0].add_run("Description").bold = True
    hdr_cells[0].paragraphs[0].add_run("\n\n\n")
    hdr_cells[0].paragraphs[0].add_run("Date Entry Created").bold = True
    hdr_cells[0].paragraphs[0].add_run("\n\n\n")
    hdr_cells[0].paragraphs[0].add_run("POC").bold = True
    hdr_cells[0].paragraphs[0].add_run("\n")
    document.add_paragraph("", style=boan_style)
    docx_dir = "C:\\Users\\ahnlab\\Desktop\\업무\\외부동향\\"
    docx_file = "일일 외부 동향_{0}.docx".format(dt.strftime(dt.now(), "%Y%m%d"))
    if os.path.isfile(docx_dir + docx_file) == True:
        ask = messagebox.askquestion("확인창", "이미 일일 외부 동향 폴더에 '{0}'파일이 있습니다.\n덮어 쓰시겠습니까?".format(docx_file))
        if ask == "yes":
            document.save(docx_dir + docx_file)
            os.popen(docx_dir + docx_file)
        else:
            messagebox.showinfo("알림창", "dump.docx파일로 워드를 실행합니다.")
            document.save("dump.docx")
            os.popen("dump.docx")
    else:
        try:
            document.save(docx_dir + docx_file)
            os.popen(docx_dir + docx_file)
        except:
            document.save("dump.docx")
            os.popen("dump.docx")

    cap_dir = "C:\\Users\\ahnlab\\Desktop\\업무\\외부동향 캡쳐\\{0}\\".format(dt.strftime(dt.now(), "%m%d"))
    if os.path.isdir(cap_dir) == True:
        ask_capt = messagebox.showerror("캡쳐 오류", "이미 '{}'경로에 폴더가 존재하여, 캡쳐를 진행하지 않습니다.".format(cap_dir))
        TEXT_for_mail.insert(END, "\n\n캡쳐를 진행하지 않았습니다.")

    elif os.path.isdir(cap_dir) == False:
        TEXT_for_mail.insert(END, "\n\n캡쳐를 시작합니다.")
        os.mkdir(cap_dir)
        Cap_count = 1
        option1 = webdriver.ChromeOptions()
        option1.add_argument('headless')
        C_URL = ['https://www.exploit-db.com/',
                 'https://www.kb.cert.org/vuls/bypublished/desc/',
                 'https://www.krcert.or.kr/data/secNoticeList.do',
                 'https://www.krcert.or.kr/data/secInfoList.do', ]
        for C in C_URL:
            driver = webdriver.Chrome(executable_path='chromedriver', options=option1)
            driver.get(url=C)
            sleep(3)
            driver.save_screenshot("C:\\Users\\ahnlab\\Desktop\\업무\\외부동향 캡쳐\\{0}\\capture{1}.png".format(
                dt.strftime(dt.now(), "%m%d"), Cap_count))
            driver.close()
            Cap_count = Cap_count + 1
        TEXT_for_mail.insert(END, "\n\nExploitDB, KRCERT(공지,취약점), USCERT의 캡쳐가 완료되었습니다.")


if __name__ == '__main__':
    LIST_Boan = []
    root = Tk()
    root.title("외부동향 보고서 작성 툴")
    root.geometry('710x660+400+400')
    root.configure(background='white')
    win = Frame(root)

    LB_for_boan = Label(root, text="보안뉴스", pady=10, bg="white")
    LB_for_boan.grid(column=0, row=0)
    Combo_Update()

    BT_for_crawl = Button(root, text='crawl', command=boannews_crawl, bg='white', height=1, bd=1, pady=10)
    BT_for_crawl.grid(column=2, row=0, sticky=W + E)

    TEXT_for_mail = Text(root, width=100, height=40, bd=3, pady=10)
    TEXT_for_mail.grid(column=0, row=1, columnspan=3)
    BT_for_docx = Button(root, text='작성', command=lambda: process(target=default_write).start(), bg='white', height=1,
                         bd=1, pady=10)
    BT_for_docx.grid(column=0, row=2, columnspan=3, sticky=W + E)
    TEXT_for_mail.insert(END, "# 사용방법 #\n1. crawl버튼 클릭 -> 뉴스 선택 -> 작성 버튼 \n\n- 보안뉴스만 선택하며 나머지는 자동작성 됨\n\n"
                              "- 프로그램 실행일 기준 -1day의 게시물을 크롤링 (00시 넘어가면 프로그램 재시작 필요)\n\n"
                              "- 재시작 방법 : Auto_Docx.py 클릭 후 'CTRL + Shift + F10' \n\n\n"
                              "# 현재 구현된 기능 #\n1.보안뉴스 : -1Day 이후 업로드된 뉴스 목록화\n- 3개의 메인뉴스는 고정이여서 주말의 경우 전전날 뉴스가 표기될 수 있음\n\n"
                              "2.보안공지 : -1day 이후 업로드된 KRCERT(권고문,취약점), USCERT 게시물(UTC 기준 -1day)\n- KRCERT의 게시물 내의 표형식은 텍스트만 표기됨\n\n"
                              "3.보안 취약점 : 금융ISAC API 추후 업데이트 시 가능\n\n"
                              "4.ExploitDB : -1day 이후로 업로드된 ExploitDB (Remote, Webapps)\n\n"
                              "5.Github : 인터넷망 PC Github 접근 불가\n\n\n"
                              "# 추가 기능 #\n1. ExploitDB, KRCERT(보안공지, 취약점정보), USCERT 자동 캡쳐 기능 추가\n\n"
                              "- 외부동향 캡쳐리스트에 날짜별 자동 저장\n\n"
                              "- 금융ISAC(외부동향,보안권고,주요위협정보) 별도 캡쳐 필요\n\n"
                              "- 이미 날짜 폴더가 있을 경우 경고 문구 추가\n\n\n"
                              "2. 일일 외부동향 docx 자동 저장 기능\n\n"
                              "- 작성 버튼 클릭 시 자동 작성 뒤 파일 Open\n\n"
                              "- 이미 파일이 있을 경우 '덮어쓰기 or Dump.docx'로 실행 가능\n\n")

    dayday = "{}".format(dt.strftime(dt.now(), "%Y.%m.%d"))
    document = docx.Document()

    main = document.styles['Normal']
    main.font.name = '맑은 고딕'
    main.font.size = Pt(12)
    rFonts = main.element.rPr.rFonts
    rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    Head = document.styles.add_style('Head', WD_STYLE_TYPE.PARAGRAPH)
    Head.font.name = '맑은 고딕'
    Head.font.size = Pt(42)
    Head.font.color.rgb = RGBColor(31, 71, 137)
    rFonts2 = Head.element.rPr.rFonts
    rFonts2.set(qn("w:eastAsia"), "맑은 고딕")

    sub_Head = document.styles.add_style('sub_Head', WD_STYLE_TYPE.PARAGRAPH)
    sub_Head.font.name = '맑은 고딕'
    sub_Head.font.size = Pt(12)
    sub_Head.font.color.rgb = RGBColor(31, 71, 137)
    rFonts3 = sub_Head.element.rPr.rFonts
    rFonts3.set(qn("w:eastAsia"), "맑은 고딕")

    sub_Head2 = document.styles.add_style('sub_Head2', WD_STYLE_TYPE.PARAGRAPH)
    sub_Head2.font.name = '맑은 고딕'
    sub_Head2.font.size = Pt(16)
    sub_Head2.font.bold = True
    sub_Head2.font.color.rgb = RGBColor(31, 71, 137)
    rFonts4 = sub_Head2.element.rPr.rFonts
    rFonts4.set(qn("w:eastAsia"), "맑은 고딕")

    boan_style = document.styles.add_style('boan_style', WD_STYLE_TYPE.PARAGRAPH)
    boan_style.font.name = '맑은 고딕'
    boan_style.font.size = Pt(14)
    boan_style.font.bold = True
    rFonts4 = boan_style.element.rPr.rFonts
    rFonts4.set(qn("w:eastAsia"), "맑은 고딕")

    source_url = document.styles.add_style('source_url', WD_STYLE_TYPE.PARAGRAPH)
    source_url.font.name = '맑은 고딕'
    source_url.font.size = Pt(11)
    source_url.font.bold = False
    rFonts4 = boan_style.element.rPr.rFonts
    rFonts4.set(qn("w:eastAsia"), "맑은 고딕")

    win.mainloop()
